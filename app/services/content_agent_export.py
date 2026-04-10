from __future__ import annotations

import re
import zipfile
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.shared import Inches


EXPORT_DIR = Path("exports") / "content-agent"
IMAGE_LINE_PATTERN = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
CAPTION_LINE_PATTERN = re.compile(r"^\*(.+)\*$")


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", (value or "").strip().lower()).strip("-")
    return slug or "content-agent-export"


def _export_path(goal: str, suffix: str) -> Path:
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return EXPORT_DIR / f"{_slugify(goal)[:60]}-{uuid4().hex[:8]}.{suffix}"


def _public_url(path: Path) -> str:
    try:
        relative = path.resolve().relative_to(Path("exports").resolve())
    except Exception:
        relative = path.name
    normalized = str(relative).replace("\\", "/")
    return f"/exports/{normalized}"


def export_markdown(goal: str, article_markdown: str) -> dict[str, str]:
    path = _export_path(goal, "md")
    path.write_text(article_markdown, encoding="utf-8")
    return {"path": str(path.resolve()), "public_url": _public_url(path), "filename": path.name}


def export_docx(goal: str, article_markdown: str) -> dict[str, str]:
    path = _export_path(goal, "docx")
    document = Document()

    lines = str(article_markdown or "").splitlines()
    for line in lines:
        trimmed = line.strip()
        if not trimmed:
            document.add_paragraph("")
            continue

        image_match = IMAGE_LINE_PATTERN.match(trimmed)
        if image_match:
            image_path = _resolve_export_image_path(image_match.group(2))
            if image_path and image_path.exists():
                document.add_picture(str(image_path), width=Inches(6.2))
            continue

        caption_match = CAPTION_LINE_PATTERN.match(trimmed)
        if caption_match:
            paragraph = document.add_paragraph()
            paragraph.add_run(caption_match.group(1)).italic = True
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", trimmed)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(heading_match.group(2), level=level)
            continue

        unordered_match = re.match(r"^[-*]\s+(.+)$", trimmed)
        if unordered_match:
            document.add_paragraph(unordered_match.group(1), style="List Bullet")
            continue

        ordered_match = re.match(r"^\d+\.\s+(.+)$", trimmed)
        if ordered_match:
            document.add_paragraph(ordered_match.group(1), style="List Number")
            continue

        meta_match = re.match(r"^(Meta Title|Meta Description)\s*:\s*(.+)$", trimmed, flags=re.IGNORECASE)
        if meta_match:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"{meta_match.group(1)}: ").bold = True
            paragraph.add_run(meta_match.group(2))
            continue

        document.add_paragraph(trimmed)

    document.save(path)
    return {"path": str(path.resolve()), "public_url": _public_url(path), "filename": path.name}


def export_images_zip(goal: str, images: list[dict[str, str]]) -> dict[str, str] | None:
    valid_images = [item for item in images if item.get("local_path")]
    if not valid_images:
        return None
    path = _export_path(goal, "zip")
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for index, image in enumerate(valid_images, start=1):
            source_path = Path(image["local_path"])
            if not source_path.exists():
                continue
            suffix = source_path.suffix or ".png"
            image_name = _slugify(image.get("title") or f"image-{index}")[:60]
            archive.write(source_path, arcname=f"{index:02d}-{image_name}{suffix}")
    return {"path": str(path.resolve()), "public_url": _public_url(path), "filename": path.name}


def _resolve_export_image_path(url: str) -> Path | None:
    value = str(url or "").strip()
    if not value:
        return None
    if value.startswith("/exports/"):
        return Path("exports") / value.removeprefix("/exports/")
    candidate = Path(value)
    return candidate if candidate.exists() else None
